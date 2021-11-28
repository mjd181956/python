from docx import Document #导入库
document=Document() #创建word文档
#加载word文档
#document=Document("")

#添加标题
document.add_heading("我是默认1级标题")
document.add_heading("我是0级标题",0)
document.add_heading("我是1级标题",1)
document.add_heading("我是2级标题",2)
document.add_heading("我是3级表题",3)

#添加段落
paragraph=document.add_paragraph("我是A段落")
paragraph=document.add_paragraph("我是B段落")

#插入上一段前
paragraph=document.insert_paragraph_from("我会将段落插入在我上一个段落上面")
paragraph=document.add_paragraph("我是C段落")

#在指定段落前插入
par=document.paragraphs[1]
par.insert_paragraph_from("hello,world")

#添加分页符
document.add_page_break()
paragraph=document.add_paragraph("因为上面有分页符，所以我们自动加入到下一页")

#保存文档
document.save('./seventeenth/word.docx')
